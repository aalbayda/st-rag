
from __future__ import annotations

import os

import pytest

os.environ.setdefault("OPENROUTER_API_KEY", "secret-or")
os.environ.setdefault("PINECONE_API_KEY", "secret-p")
os.environ.setdefault("PINECONE_INDEX_NAME", "rag-dense")


@pytest.fixture(autouse=True)
def clear_settings_and_client_caches():
    from app.config import get_settings

    get_settings.cache_clear()

    try:
        from app.clients import openai_client

        openai_client.get_openai_client.cache_clear()
    except (ImportError, AttributeError):
        pass

    try:
        from app.clients import pinecone_client

        pinecone_client.get_pinecone.cache_clear()
    except (ImportError, AttributeError):
        pass

    yield

    from app.config import get_settings as _gs

    _gs.cache_clear()

    try:
        from app.clients import openai_client

        openai_client.get_openai_client.cache_clear()
    except (ImportError, AttributeError):
        pass

    try:
        from app.clients import pinecone_client

        pinecone_client.get_pinecone.cache_clear()
    except (ImportError, AttributeError):
        pass


PAGE_ONE_MARKER = "ALPHA_PAGE_ONE_MARKER"
PAGE_TWO_MARKER = "BETA_PAGE_TWO_MARKER"


@pytest.fixture(scope="session")
def sample_pdf_path(tmp_path_factory):
    import pymupdf

    tmp_dir = tmp_path_factory.mktemp("fixtures")
    pdf_path = tmp_dir / "sample.pdf"

    doc = pymupdf.open()

    page1 = doc.new_page(width=612, height=792)
    page1.insert_text(
        (72, 72),
        f"Page one content. {PAGE_ONE_MARKER} This is the first page.",
        fontsize=12,
    )

    page2 = doc.new_page(width=612, height=792)
    page2.insert_text(
        (72, 72),
        f"Page two content. {PAGE_TWO_MARKER} This is the second page.",
        fontsize=12,
    )

    doc.save(str(pdf_path))
    doc.close()

    return str(pdf_path)


@pytest.fixture(scope="session")
def sample_docx_path(tmp_path_factory):
    from docx import Document

    tmp_dir = tmp_path_factory.mktemp("fixtures_docx")
    docx_path = tmp_dir / "sample.docx"

    doc = Document()
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph("First paragraph under Introduction.")
    doc.add_paragraph("Second paragraph under Introduction.")

    doc.save(str(docx_path))

    return str(docx_path)
